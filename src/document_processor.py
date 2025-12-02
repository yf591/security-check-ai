"""
ドキュメント処理モジュール
PDF、Word、Excel、テキストファイルからQ&Aペアを抽出
"""

import os
import re
from typing import List, Dict, Optional
from pathlib import Path
import logging

import pandas as pd
from pypdf import PdfReader
from docx import Document

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class DocumentProcessor:
    """各種ドキュメントフォーマットからQ&Aペアを抽出"""

    def __init__(self):
        self.supported_extensions = [".pdf", ".docx", ".xlsx", ".xls", ".txt", ".csv"]

    def process_file(self, file_path: str) -> List[Dict[str, str]]:
        """
        ファイルを処理してQ&Aペアのリストを返す

        Args:
            file_path: 処理するファイルのパス

        Returns:
            [{"question": "質問文", "answer": "回答文", "source": "ファイル名"}, ...]
        """
        file_path = Path(file_path)
        extension = file_path.suffix.lower()

        if extension not in self.supported_extensions:
            logger.warning(f"非対応のファイル形式: {extension}")
            return []

        logger.info(f"処理中: {file_path.name}")

        try:
            if extension == ".pdf":
                return self._process_pdf(file_path)
            elif extension == ".docx":
                return self._process_docx(file_path)
            elif extension in [".xlsx", ".xls"]:
                return self._process_excel(file_path)
            elif extension == ".txt":
                return self._process_text(file_path)
            elif extension == ".csv":
                return self._process_csv(file_path)
        except Exception as e:
            logger.error(f"ファイル処理エラー {file_path.name}: {str(e)}")
            return []

    def _process_pdf(self, file_path: Path) -> List[Dict[str, str]]:
        """PDFファイルからテキストを抽出してQ&Aペアに変換"""
        qa_pairs = []

        reader = PdfReader(str(file_path))
        full_text = ""

        for page in reader.pages:
            full_text += page.extract_text() + "\n"

        # Q&Aペアの抽出を試みる
        qa_pairs = self._extract_qa_pairs(full_text, file_path.name)

        return qa_pairs

    def _process_docx(self, file_path: Path) -> List[Dict[str, str]]:
        """Wordファイルからテキストを抽出してQ&Aペアに変換"""
        doc = Document(str(file_path))
        full_text = "\n".join([paragraph.text for paragraph in doc.paragraphs])

        # テーブルも処理
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join([cell.text for cell in row.cells])
                full_text += "\n" + row_text

        qa_pairs = self._extract_qa_pairs(full_text, file_path.name)

        return qa_pairs

    def _process_excel(self, file_path: Path) -> List[Dict[str, str]]:
        """Excelファイルからデータを読み込んでQ&Aペアに変換"""
        qa_pairs = []

        # 全シートを読み込み
        excel_file = pd.ExcelFile(file_path)

        for sheet_name in excel_file.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)

            # 列名に「質問」「Q」などが含まれているか確認
            question_col = None
            answer_col = None

            for col in df.columns:
                col_lower = str(col).lower()
                if any(
                    keyword in col_lower
                    for keyword in ["質問", "q", "question", "query"]
                ):
                    question_col = col
                elif any(
                    keyword in col_lower
                    for keyword in ["回答", "a", "answer", "response"]
                ):
                    answer_col = col

            # Q&A形式のデータがある場合
            if question_col is not None and answer_col is not None:
                for _, row in df.iterrows():
                    question = str(row[question_col]).strip()
                    answer = str(row[answer_col]).strip()

                    if question and answer and question != "nan" and answer != "nan":
                        qa_pairs.append(
                            {
                                "question": question,
                                "answer": answer,
                                "source": f"{file_path.name} - {sheet_name}",
                            }
                        )
            else:
                # Q&A形式でない場合、全テキストを結合して処理
                text = df.to_string()
                qa_pairs.extend(
                    self._extract_qa_pairs(text, f"{file_path.name} - {sheet_name}")
                )

        return qa_pairs

    def _process_text(self, file_path: Path) -> List[Dict[str, str]]:
        """テキストファイルからQ&Aペアを抽出"""
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            content = f.read()

        return self._extract_qa_pairs(content, file_path.name)

    def _process_csv(self, file_path: Path) -> List[Dict[str, str]]:
        """CSVファイルからQ&Aペアを抽出"""
        qa_pairs = []

        try:
            df = pd.read_csv(file_path, encoding="utf-8")
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding="cp932")

        # Q&A列を探す
        question_col = None
        answer_col = None

        for col in df.columns:
            col_lower = str(col).lower()
            if any(
                keyword in col_lower for keyword in ["質問", "q", "question", "query"]
            ):
                question_col = col
            elif any(
                keyword in col_lower for keyword in ["回答", "a", "answer", "response"]
            ):
                answer_col = col

        if question_col is not None and answer_col is not None:
            for _, row in df.iterrows():
                question = str(row[question_col]).strip()
                answer = str(row[answer_col]).strip()

                if question and answer and question != "nan" and answer != "nan":
                    qa_pairs.append(
                        {
                            "question": question,
                            "answer": answer,
                            "source": file_path.name,
                        }
                    )
        else:
            text = df.to_string()
            qa_pairs.extend(self._extract_qa_pairs(text, file_path.name))

        return qa_pairs

    def _extract_qa_pairs(self, text: str, source: str) -> List[Dict[str, str]]:
        """
        テキストからQ&Aペアを抽出

        パターン:
        - Q: ... A: ...
        - Q. ... A. ...
        - 質問: ... 回答: ...
        - 【質問】... 【回答】...
        """
        qa_pairs = []

        # 複数のパターンでQ&Aペアを抽出
        patterns = [
            r"[Qq][.:]\s*(.+?)\s*[Aa][.:]\s*(.+?)(?=[Qq][.:]|$)",
            r"質問[.::\s](.+?)\s*回答[.::\s](.+?)(?=質問|$)",
            r"【質問】\s*(.+?)\s*【回答】\s*(.+?)(?=【質問】|$)",
            r"Q[0-9]+[.:]\s*(.+?)\s*A[0-9]+[.:]\s*(.+?)(?=Q[0-9]+|$)",
        ]

        for pattern in patterns:
            matches = re.findall(pattern, text, re.DOTALL | re.IGNORECASE)
            for match in matches:
                if len(match) == 2:
                    question = match[0].strip()
                    answer = match[1].strip()

                    if len(question) > 5 and len(answer) > 5:  # 最小長チェック
                        qa_pairs.append(
                            {"question": question, "answer": answer, "source": source}
                        )

        # パターンマッチで見つからない場合、段落単位でチャンクとして保存
        if not qa_pairs:
            paragraphs = [p.strip() for p in text.split("\n\n") if len(p.strip()) > 50]
            for i, paragraph in enumerate(paragraphs):
                qa_pairs.append(
                    {
                        "question": f"{source} - セクション {i+1}",
                        "answer": paragraph,
                        "source": source,
                    }
                )

        return qa_pairs


def process_directory(directory_path: str) -> List[Dict[str, str]]:
    """
    ディレクトリ内の全ファイルを処理

    Args:
        directory_path: 処理するディレクトリのパス

    Returns:
        全ファイルから抽出されたQ&Aペアのリスト
    """
    processor = DocumentProcessor()
    all_qa_pairs = []

    directory = Path(directory_path)

    if not directory.exists():
        logger.error(f"ディレクトリが存在しません: {directory_path}")
        return []

    for file_path in directory.rglob("*"):
        if (
            file_path.is_file()
            and file_path.suffix.lower() in processor.supported_extensions
        ):
            qa_pairs = processor.process_file(str(file_path))
            all_qa_pairs.extend(qa_pairs)
            logger.info(f"{file_path.name} から {len(qa_pairs)} 件のQ&Aペアを抽出")

    logger.info(f"合計 {len(all_qa_pairs)} 件のQ&Aペアを抽出しました")

    return all_qa_pairs


if __name__ == "__main__":
    # テスト用
    test_dir = "./data/raw"
    qa_pairs = process_directory(test_dir)
    print(f"抽出されたQ&Aペア: {len(qa_pairs)}件")

    if qa_pairs:
        print("\n最初の3件:")
        for qa in qa_pairs[:3]:
            print(f"\n質問: {qa['question'][:100]}...")
            print(f"回答: {qa['answer'][:100]}...")
            print(f"出典: {qa['source']}")
